
import pytesseract
from PIL import Image
import docx

text = ""

image_path = f"1.jpg" #image to extract text from
image = Image.open(image_path)
pytesseract.pytesseract.tesseract_cmd = 'D:\\program_files\\Tesseract-OCR\\tesseract.exe'
text += pytesseract.image_to_string(image, lang="heb")

document = docx.Document()

document.add_heading('', 0)
document.add_paragraph(text)
document.save('output.docx')
print("images converted successfully")
